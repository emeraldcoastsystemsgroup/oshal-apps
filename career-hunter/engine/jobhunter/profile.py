# CHANGE LOG
# -----------------------------------------------------------------------------
# SEQ | AUTHOR                                    | DESCRIPTION
# -----------------------------------------------------------------------------
# 1 | maintainer@emeraldcoastsystemsgroup.com | Add the Resume Studio MASTER document mapping: base_document() (straight, model-free profile -> editor shape) and replace_resume_fields() (whitelist write-back reusing augment's backup/audit/atomic-replace rails), so the durable profile every tailored resume is generated from finally has an editor.

"""Loads the user's career_db.json and renders a compact profile for prompts."""
from __future__ import annotations
import json
import logging
import os
import stat
from datetime import datetime
from functools import lru_cache
from pathlib import Path

from . import config


logger = logging.getLogger(__name__)

_PROFILE_MAX_BYTES = 4 * 1024 * 1024
_PROFILE_BACKUP_LIMIT = 20
_ENRICHMENT_LOG_MAX_BYTES = 1024 * 1024
_AUGMENT_FACT_MAX_CHARS = 16_000
_AUGMENT_FACT_MAX_BYTES = 16 * 1024
_AUGMENT_PROMPT_MAX_BYTES = 64 * 1024
_AUGMENT_PROFILE_TRUNCATION = "\n[CURRENT PROFILE OUTLINE TRUNCATED TO INPUT BUDGET]\n"


def _read_profile_text() -> str:
    """Read one regular profile only after enforcing the same byte ceiling as writes."""
    metadata = config.CAREER_DB.lstat()
    if not stat.S_ISREG(metadata.st_mode):
        raise ValueError("career profile must be a regular file")
    if metadata.st_size > _PROFILE_MAX_BYTES:
        raise ValueError("career profile exceeds the storage limit")
    return config.CAREER_DB.read_text(encoding="utf-8")


def _render_profile(raw: dict) -> str:
    """Serialize one profile and reject model output beyond the durable store budget."""
    rendered = json.dumps(raw, indent=2, ensure_ascii=False)
    if len(rendered.encode("utf-8")) > _PROFILE_MAX_BYTES:
        raise ValueError("career profile exceeds the storage limit")
    return rendered


def _replace_profile(rendered: str) -> None:
    """Atomically publish already-bounded profile bytes in the caller-owned directory."""
    temporary = config.CAREER_DB.with_suffix(".json.tmp")
    temporary.write_text(rendered, encoding="utf-8")
    os.replace(temporary, config.CAREER_DB)


@lru_cache(maxsize=1)
def load() -> dict:
    if not config.CAREER_DB.exists():
        return {}
    return json.loads(_read_profile_text())


def summary(max_chars: int = 3500, include_oshal: bool = False) -> str:
    """A dense plain-text profile for AI scoring/generation prompts.

    include_oshal=True appends the user's independent open-source / R&D work (the
    `oss_portfolio` block) so the scorer credits their
    hands-on agentic / LLM / MCP / RAG depth. This is used for SCORING (private,
    never shown to an employer); resume/cover generation gates it separately. The
    OSHAL block is appended AFTER the base truncation so it always survives.
    """
    d = load()
    if not d:
        return "No career profile found."
    p = d.get("profile", {})
    lines = [f"CANDIDATE: {p.get('name','')} ({p.get('credential','')}) — {p.get('location','')}",
             f"Clearance: {p.get('clearance','')}; {p.get('citizenship','')}",
             f"Summary: {p.get('experience_summary','')}", ""]
    lines.append("EXPERIENCE:")
    for r in d.get("roles", [])[:12]:
        span = f"{r.get('start','')}–{r.get('end') or 'present'}"
        lines.append(f"- {r.get('title','')} @ {r.get('org','')} ({span})")
        for b in (r.get("deliverables") or [])[:3]:
            lines.append(f"    • {b}")
    skills = d.get("skills", {})
    flat = []
    for g in skills.values():
        flat.extend(g.get("items", []))
    lines.append("\nSKILLS: " + ", ".join(flat[:60]))
    lines.append("\nMETRICS: " + " | ".join(d.get("metrics_bank", [])[:8]))
    certs = d.get("certifications", [])
    if certs:
        lines.append("\nCERTS: " + ", ".join(certs))
    text = "\n".join(lines)[:max_chars]
    if include_oshal:
        oss = d.get("oss_portfolio") or {}
        blurb = oss.get("scoring_blurb")
        if blurb:
            extra = "\n\n" + blurb
            oss_skills = oss.get("skills") or []
            if oss_skills:
                extra += "\nOSS SKILLS: " + ", ".join(oss_skills)
            text += extra
    return text


_ENH_SYS = (
    "You maintain a candidate's structured career database. You merge NEW, candidate-confirmed "
    "TRUE facts into it NON-DESTRUCTIVELY and conservatively. Never invent facts beyond what the "
    "candidate states, never remove or weaken existing content, never exaggerate. Prefer adding to "
    "an EXISTING skill group or role rather than creating new ones. Return ONLY a strict JSON patch "
    "of ADDITIONS.")


def _low(xs):
    return {str(x).strip().lower() for x in xs}


def _augmentation_sections(raw: dict, facts: str) -> tuple[str, str]:
    """Build separate profile and trusted-fact sections so profile text cannot spoof a delimiter."""
    groups = {g: (raw.get("skills", {}).get(g, {}) or {}).get("items", []) for g in raw.get("skills", {})}
    roles = raw.get("roles", [])
    outline = (
        "CURRENT SKILL GROUPS (name: sample items):\n"
        + "\n".join(f"- {g}: {', '.join(items[:8])}" for g, items in groups.items())
        + "\n\nROLES (index | title @ org):\n"
        + "\n".join(f"{i} | {r.get('title','')} @ {r.get('org','')}" for i, r in enumerate(roles))
    )
    tail = (
        "\n\nNEW CANDIDATE-CONFIRMED FACTS TO MERGE:\n" + facts
        + "\n\nReturn STRICT JSON only:\n{\n"
        '  "skills_add": {"<existing group name, or a concise new group>": ["skill phrase", ...]},\n'
        '  "role_bullets_add": [{"role_index": <int from the list above — the BEST-matching role>, '
        '"bullets": ["achievement bullet, active voice, quantified only if the facts support it"]}],\n'
        '  "metrics_add": ["quantified metric phrase, only if the facts include one"],\n'
        '  "changelog": ["short human-readable line per addition"]\n}\n'
        "Only include additions clearly supported by the facts; do not duplicate existing content."
    )
    return outline, tail


def _utf8_prefix(value: str, limit: int) -> tuple[str, bool, int]:
    """Return a deterministic valid UTF-8 prefix plus truncation and original-byte diagnostics."""
    encoded = value.encode("utf-8", errors="replace")
    if len(encoded) <= limit:
        return value, False, len(encoded)
    return encoded[:limit].decode("utf-8", errors="ignore"), True, len(encoded)


def _bounded_augmentation_context(raw: dict, supplied_facts: str) -> tuple[str, str, dict]:
    """Bound combined model text to 64 KiB and report deterministic UTF-8 byte diagnostics."""
    char_bounded = supplied_facts[:_AUGMENT_FACT_MAX_CHARS]
    facts, byte_cut, _ = _utf8_prefix(char_bounded, _AUGMENT_FACT_MAX_BYTES)
    outline, tail = _augmentation_sections(raw, facts)
    system_bytes = len(_ENH_SYS.encode("utf-8"))
    context_budget = _AUGMENT_PROMPT_MAX_BYTES - system_bytes
    tail_bytes = len(tail.encode("utf-8"))
    if tail_bytes > context_budget:
        raise ValueError("augmentation facts and instructions exceed the model input limit")
    available = context_budget - tail_bytes
    bounded_outline, outline_cut, outline_input_bytes = _utf8_prefix(outline, available)
    if outline_cut:
        marker_bytes = len(_AUGMENT_PROFILE_TRUNCATION.encode("utf-8"))
        marker = _AUGMENT_PROFILE_TRUNCATION if marker_bytes <= available else ""
        bounded_outline, _, _ = _utf8_prefix(outline, available - len(marker.encode("utf-8")))
        bounded_outline += marker
    context = bounded_outline + tail
    model_input_bytes = system_bytes + len(context.encode("utf-8"))
    if model_input_bytes > _AUGMENT_PROMPT_MAX_BYTES:
        raise ValueError("augmentation model input exceeds the byte limit")
    supplied_bytes = len(supplied_facts.encode("utf-8", errors="replace"))
    diagnostics = {
        "prompt_limit_bytes": _AUGMENT_PROMPT_MAX_BYTES,
        "model_input_bytes": model_input_bytes,
        "profile_input_bytes": outline_input_bytes,
        "profile_bytes": len(bounded_outline.encode("utf-8")),
        "profile_truncated": outline_cut,
        "facts_input_bytes": supplied_bytes,
        "facts_bytes": len(facts.encode("utf-8")),
        "facts_truncated": len(supplied_facts) > _AUGMENT_FACT_MAX_CHARS or byte_cut,
    }
    return context, facts, diagnostics


def _augmentation_result(changed: bool, changelog: list, diagnostics: dict | None = None) -> dict:
    """Expose successful execution separately from whether additive persistence changed bytes."""
    result = {"ok": True, "changed": changed, "changelog": changelog}
    if diagnostics is not None:
        result["diagnostics"] = diagnostics
    return result


def _append_unique(target: list, additions) -> bool:
    """Append distinct non-empty strings without changing existing order or content."""
    have = _low(target)
    changed = False
    for item in additions or []:
        value = item.strip() if isinstance(item, str) else ""
        if value and value.lower() not in have:
            target.append(value)
            have.add(value.lower())
            changed = True
    return changed


def _apply_augmentation(raw: dict, patch: dict) -> bool:
    """Apply only additive skills, role bullets, and metrics from a validated patch shape."""
    changed = False
    skills = raw.setdefault("skills", {})
    for group, items in (patch.get("skills_add") or {}).items():
        if not items:
            continue
        grp = skills.setdefault(group, {"items": []})
        grp.setdefault("items", [])
        changed = _append_unique(grp["items"], items) or changed
    roles = raw.get("roles", [])
    for role_patch in (patch.get("role_bullets_add") or []):
        idx, bullets = role_patch.get("role_index"), (role_patch.get("bullets") or [])
        if isinstance(idx, int) and 0 <= idx < len(roles) and bullets:
            changed = _append_unique(roles[idx].setdefault("deliverables", []), bullets) or changed
    return _append_unique(raw.setdefault("metrics_bank", []), patch.get("metrics_add")) or changed


def _rotate_profile_backups(directory: Path) -> None:
    """Retain only the newest bounded set of full-profile recovery snapshots."""
    backups = sorted(
        directory.glob("career_db.*.json"),
        key=lambda item: item.stat().st_mtime_ns,
        reverse=True,
    )
    for backup in backups[_PROFILE_BACKUP_LIMIT:]:
        backup.unlink(missing_ok=True)


def _bounded_audit_record(stamp: str, facts: str, changelog: list) -> bytes:
    """Serialize one bounded enrichment record without retaining unbounded model output."""
    safe_log = [str(item)[:500] for item in changelog[:50]]
    record = {"at": stamp, "facts": facts[:_AUGMENT_FACT_MAX_CHARS], "changelog": safe_log}
    return (json.dumps(record, ensure_ascii=False) + "\n").encode("utf-8")


def _write_bounded_audit(stamp: str, facts: str, changelog: list) -> None:
    """Atomically retain only a complete-line tail within the enrichment audit byte budget."""
    audit = config.CAREER_DB.parent / "enrichment_log.jsonl"
    record = _bounded_audit_record(stamp, facts, changelog)
    keep = max(0, _ENRICHMENT_LOG_MAX_BYTES - len(record))
    prior = b""
    if keep and audit.exists():
        with audit.open("rb") as source:
            offset = max(0, audit.stat().st_size - keep)
            source.seek(offset)
            prior = source.read(keep)
        if offset:
            prior = prior.partition(b"\n")[2]
    temporary = audit.with_suffix(".jsonl.tmp")
    temporary.write_bytes(prior + record[-_ENRICHMENT_LOG_MAX_BYTES:])
    os.replace(temporary, audit)


def _persist_augmentation(text: str, raw: dict, facts: str, changelog: list) -> None:
    """Bound backups/audit, atomically replace the profile, and invalidate its cache."""
    rendered = _render_profile(raw)
    bdir = config.CAREER_DB.parent / "backups"
    bdir.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S-%f")
    (bdir / f"career_db.{stamp}.json").write_text(text, encoding="utf-8")
    _rotate_profile_backups(bdir)
    _replace_profile(rendered)

    # audit log + drop the cache so the next load()/generate sees the new data
    try:
        _write_bounded_audit(stamp, facts, changelog)
    except Exception:
        logger.exception("career profile enrichment audit write failed")
    load.cache_clear()


def augment(facts: str) -> dict:
    """Merge candidate-confirmed facts additively, with backup, audit, and atomic replacement."""
    from . import enrich
    supplied_facts = str(facts or "").strip()
    if not supplied_facts:
        return _augmentation_result(False, [])
    text = _read_profile_text()
    raw = json.loads(text)
    context, accepted_facts, diagnostics = _bounded_augmentation_context(raw, supplied_facts)
    patch = enrich.parse_json(enrich.complete(_ENH_SYS, context, max_tokens=1500)) or {}
    reported = patch.get("changelog") or []
    changelog = list(reported) if isinstance(reported, list) else []
    if not _apply_augmentation(raw, patch):
        return _augmentation_result(False, [], diagnostics)
    _persist_augmentation(text, raw, accepted_facts, changelog)
    return _augmentation_result(True, changelog, diagnostics)


# ── Resume ingest — build a NEW user's career_db.json from their uploaded resume ──
_INGEST_SYS = (
    "You convert a candidate's resume into a STRICT, structured career database as JSON. "
    "Use ONLY facts stated in the resume — never invent employers, dates, titles, metrics, "
    "clearances, or certifications. If the resume does not state something, leave it empty or omit it. "
    "Do not exaggerate. Return JSON only, no prose.")


def has_profile() -> bool:
    """True once the user has an indexed resume (a non-trivial career_db.json)."""
    try:
        d = load()
        return bool(d.get("roles") or (d.get("profile", {}) or {}).get("experience_summary"))
    except Exception:
        return False


_EXTRACT_MAX_CHARS = 32_000
_ARCHIVE_MAX_FILE_BYTES = 64 * 1024 * 1024
_PDF_MAX_PAGES = 100
_PDF_MAX_FILE_BYTES = 32 * 1024 * 1024
_PDF_MAX_STREAM_BYTES = 8 * 1024 * 1024
_PDF_MAX_TOTAL_DECODED_BYTES = 16 * 1024 * 1024
_DOCX_MAX_MEMBERS = 1_000
_DOCX_MAX_PARAGRAPHS = 5_000
_DOCX_MAX_ENTRY_BYTES = 12 * 1024 * 1024
_DOCX_MAX_TOTAL_BYTES = 32 * 1024 * 1024
_DOCX_MAX_COMPRESSION_RATIO = 200


def _validate_file_size(file_path: str, limit: int, label: str) -> None:
    """Bound compressed/raw input before a parser materializes its index or object graph."""
    size = Path(file_path).stat().st_size
    if size > limit:
        raise ValueError(f"{label} exceeds the input-byte limit")


def _bounded_text(parts) -> str:
    """Join an iterator without retaining more text than the downstream prompt can consume."""
    output = []
    remaining = _EXTRACT_MAX_CHARS
    for value in parts:
        separator = 1 if output else 0
        if remaining <= separator:
            break
        text = str(value or "")
        piece = text[:remaining - separator]
        output.append(piece)
        remaining -= separator + len(piece)
    return "\n".join(output)


def _validate_docx_archive(file_path: str) -> None:
    """Reject DOCX zip expansion shapes before python-docx materializes its XML/package parts."""
    import zipfile
    _validate_file_size(file_path, _ARCHIVE_MAX_FILE_BYTES, "DOCX")
    with zipfile.ZipFile(file_path) as archive:
        members = [member for member in archive.infolist() if not member.is_dir()]
        if len(members) > _DOCX_MAX_MEMBERS:
            raise ValueError("DOCX contains too many package members")
        if sum(member.file_size for member in members) > _DOCX_MAX_TOTAL_BYTES:
            raise ValueError("DOCX exceeds the uncompressed-byte limit")
        for member in members:
            if member.flag_bits & 0x1:
                raise ValueError("encrypted DOCX members are not supported")
            if member.file_size > _DOCX_MAX_ENTRY_BYTES:
                raise ValueError("DOCX member exceeds the uncompressed-byte limit")
            if member.file_size / max(1, member.compress_size) > _DOCX_MAX_COMPRESSION_RATIO:
                raise ValueError("DOCX member has a suspicious compression ratio")


def _install_pdf_decode_budget(decoded_stream_type):
    """Count each decoded pypdf stream once and reject aggregate expansion across the document."""
    original = decoded_stream_type.get_data
    seen = set()
    total = 0

    def bounded_get_data(stream):
        nonlocal total
        data = original(stream)
        marker = id(stream)
        if marker not in seen:
            seen.add(marker)
            total += len(data)
            if total > _PDF_MAX_TOTAL_DECODED_BYTES:
                raise ValueError("PDF exceeds the aggregate decoded-stream limit")
        return data

    decoded_stream_type.get_data = bounded_get_data
    return original


def _extract_pdf_text(file_path: str) -> str:
    """Extract bounded text under per-stream, aggregate-decoded, input-byte, and page limits."""
    _validate_file_size(file_path, _PDF_MAX_FILE_BYTES, "PDF")
    from pypdf import PdfReader, filters
    from pypdf.generic import DecodedStreamObject
    for name in ("ZLIB_MAX_OUTPUT_LENGTH", "LZW_MAX_OUTPUT_LENGTH", "RUN_LENGTH_MAX_OUTPUT_LENGTH",
                 "JBIG2_MAX_OUTPUT_LENGTH", "MAX_ARRAY_BASED_STREAM_OUTPUT_LENGTH",
                 "MAX_DECLARED_STREAM_LENGTH", "FLATE_MAX_BUFFER_SIZE"):
        if hasattr(filters, name):
            setattr(filters, name, _PDF_MAX_STREAM_BYTES)
    original_get_data = _install_pdf_decode_budget(DecodedStreamObject)
    try:
        reader = PdfReader(file_path)
        if len(reader.pages) > _PDF_MAX_PAGES:
            raise ValueError("PDF exceeds the page limit")
        return _bounded_text((page.extract_text() or "") for page in reader.pages)
    finally:
        DecodedStreamObject.get_data = original_get_data


def _extract_resume_text(file_path: str) -> str:
    """Pull bounded plain text from a PDF/DOCX/text resume without archive expansion hazards."""
    p = file_path.lower()
    if p.endswith(".pdf"):
        return _extract_pdf_text(file_path)
    if p.endswith(".docx"):
        import docx  # python-docx (pure-python, uses lxml already installed)
        _validate_docx_archive(file_path)
        document = docx.Document(file_path)
        if len(document.paragraphs) > _DOCX_MAX_PARAGRAPHS:
            raise ValueError("DOCX exceeds the paragraph limit")
        return _bounded_text(paragraph.text for paragraph in document.paragraphs)
    with Path(file_path).open("r", encoding="utf-8", errors="ignore") as stream:
        return stream.read(_EXTRACT_MAX_CHARS)


def build_from_resume(file_path: str) -> dict:
    """Parse an uploaded resume and write THIS user's career_db.json (the profile the
    scorer/generator read). Multi-tenant: config.CAREER_DB is the per-user path set by
    the oshal-jobhunter shim, so this never touches another user's profile.
    Returns {ok, roles, skills} or {ok:False, error}."""
    from . import enrich
    try:
        text = (_extract_resume_text(file_path) or "").strip()
    except Exception as exc:
        return {"ok": False, "error": f"could not read resume: {exc}"}
    if len(text) < 40:
        return {"ok": False, "error": "no readable text in the resume"}
    text = text[:24000]

    prompt = (
        "RESUME:\n" + text + "\n\n"
        "Return STRICT JSON only, exactly this shape:\n"
        "{\n"
        '  "profile": {"name":"","credential":"","location":"","clearance":"","citizenship":"","experience_summary":""},\n'
        '  "roles": [{"title":"","org":"","start":"YYYY-MM","end":"YYYY-MM or null","deliverables":["achievement bullet, active voice"]}],\n'
        '  "skills": {"<concise group name>": {"items": ["skill phrase"]}},\n'
        '  "metrics_bank": ["quantified achievement — only if the resume states a number"],\n'
        '  "certifications": ["certification name"]\n'
        "}\n"
        "experience_summary = a 2-3 sentence professional summary. Use only facts in the resume."
    )
    data = enrich.parse_json(enrich.complete(_INGEST_SYS, prompt, max_tokens=4000)) or {}
    if not isinstance(data, dict) or not (data.get("roles") or data.get("profile")):
        return {"ok": False, "error": "could not parse a profile from the resume"}

    data.setdefault("profile", {})
    data.setdefault("roles", [])
    data.setdefault("skills", {})
    data.setdefault("metrics_bank", [])
    data.setdefault("certifications", [])

    config.CAREER_DB.parent.mkdir(parents=True, exist_ok=True)
    try:
        _replace_profile(_render_profile(data))
    except ValueError as exc:
        return {"ok": False, "error": str(exc)}
    load.cache_clear()
    return {
        "ok": True,
        "roles": len(data.get("roles", [])),
        "skills": sum(len((g or {}).get("items", [])) for g in data.get("skills", {}).values()),
    }


# ── Absorb an arbitrary career artifact (email, LinkedIn export, status report, work sample) ──
# The conversational career agent lets a user add MORE than a resume: it extracts TRUE facts from
# the artifact and merges them via augment() (non-destructive, backed up, audited). This is the
# non-resume sibling of build_from_resume — it never rebuilds/overwrites the profile.
_ABSORB_SYS = (
    "You extract ONLY TRUE, verifiable career facts about the candidate from an uploaded document. "
    "List concrete facts they can stand behind: roles/titles/employers/dates, skills and tools they "
    "actually used, responsibilities, and quantified achievements STATED in the document. Never invent, "
    "never infer beyond what is written, never add aspirational or third-party facts. If the document "
    "has nothing about the candidate's own career, return an empty list. Output a plain bullet list.")

_ARTIFACT_KINDS = {"resume-extra", "linkedin-export", "email", "status-report", "work-sample", "other"}
_ZIP_RELEVANT = ("position", "profile", "skill", "education", "project", "certification", "experience")
_ZIP_MAX_MEMBERS = 1_000
_ZIP_MAX_RELEVANT_MEMBERS = 50
_ZIP_MAX_ENTRY_BYTES = 2 * 1024 * 1024
_ZIP_MAX_TOTAL_BYTES = 5 * 1024 * 1024
_ZIP_MAX_COMPRESSION_RATIO = 100


def _safe_zip_members(archive) -> list:
    """Validate central-directory sizes before decompressing any relevant LinkedIn export member."""
    members = archive.infolist()
    if len(members) > _ZIP_MAX_MEMBERS:
        raise ValueError("zip contains too many members")
    selected = [m for m in members if not m.is_dir()
                and m.filename.lower().endswith((".csv", ".txt"))
                and any(key in m.filename.lower() for key in _ZIP_RELEVANT)]
    if len(selected) > _ZIP_MAX_RELEVANT_MEMBERS:
        raise ValueError("zip contains too many career-data members")
    if sum(m.file_size for m in selected) > _ZIP_MAX_TOTAL_BYTES:
        raise ValueError("zip career data exceeds the uncompressed-byte limit")
    for member in selected:
        if member.flag_bits & 0x1:
            raise ValueError("encrypted zip members are not supported")
        if member.file_size > _ZIP_MAX_ENTRY_BYTES:
            raise ValueError("zip member exceeds the uncompressed-byte limit")
        if member.file_size / max(1, member.compress_size) > _ZIP_MAX_COMPRESSION_RATIO:
            raise ValueError("zip member has a suspicious compression ratio")
    return selected


def _read_bounded_zip_member(archive, member) -> str:
    """Stream one already-validated member and enforce its declared size while decompressing."""
    chunks = []
    total = 0
    with archive.open(member) as stream:
        while True:
            chunk = stream.read(min(64 * 1024, _ZIP_MAX_ENTRY_BYTES - total + 1))
            if not chunk:
                break
            total += len(chunk)
            if total > _ZIP_MAX_ENTRY_BYTES or total > member.file_size:
                raise ValueError("zip member expanded beyond its declared safe size")
            chunks.append(chunk)
    return b"".join(chunks).decode("utf-8", "ignore")


def _extract_artifact_text(file_path: str) -> str:
    """Plain text from a career artifact — reuses the resume extractor, adds csv/eml/html/json and
    the LinkedIn 'Download your data' zip (Positions/Profile/Skills/... members)."""
    p = file_path.lower()
    if p.endswith((".pdf", ".docx", ".txt", ".md")):
        return _extract_resume_text(file_path)
    if p.endswith(".zip"):
        import zipfile
        out = []
        _validate_file_size(file_path, _ARCHIVE_MAX_FILE_BYTES, "zip")
        with zipfile.ZipFile(file_path) as archive:
            for member in _safe_zip_members(archive):
                out.append(f"### {member.filename}\n" + _read_bounded_zip_member(archive, member))
        return "\n\n".join(out)
    # csv / tsv / html / eml / json / log and anything else — read as text
    with Path(file_path).open("r", encoding="utf-8", errors="ignore") as stream:
        return stream.read(_EXTRACT_MAX_CHARS)


def absorb(file_path: str, kind: str = "other") -> dict:
    """Absorb an uploaded career artifact into the profile: extract TRUE facts, then merge via
    augment() (add-only, backed up, audited to enrichment_log.jsonl). Requires an existing profile
    (upload a resume first). Returns {ok, kind, changelog} or {ok:False, error}."""
    from . import enrich
    if not has_profile():
        return {"ok": False, "error": "no career profile yet — upload a resume first, then add artifacts"}
    kind = kind if kind in _ARTIFACT_KINDS else "other"
    try:
        text = (_extract_artifact_text(file_path) or "").strip()
    except Exception as exc:
        return {"ok": False, "error": f"could not read artifact: {exc}"}
    if len(text) < 20:
        return {"ok": False, "error": "no readable text in the artifact"}
    prompt = (f"DOCUMENT KIND: {kind}\n\nDOCUMENT:\n{text[:24000]}\n\n"
              "List the candidate's TRUE career facts from this document (one per line).")
    facts = (enrich.complete(_ABSORB_SYS, prompt, max_tokens=1200) or "").strip()
    if len(facts) < 10:
        return {"ok": True, "changed": False, "kind": kind, "changelog": [],
                "note": "no new career facts found"}
    res = augment(facts)
    return {"ok": res.get("ok") is True, "changed": res.get("changed") is True,
            "kind": kind, "changelog": res.get("changelog", [])}


# ── Master-resume editor mapping (Resume Studio "master" document) ──
# The studio's per-job packets are generated FROM this durable profile; these functions give the
# surface a first-class editor over the profile itself. base_document() is a STRAIGHT mapping —
# no model call, no fabrication. replace_resume_fields() is the whitelist write-back that reuses
# augment()'s safety rails (bounded sizes, backup rotation, audit record, atomic replace).
_MASTER_TEXT_FIELDS = (
    ("headline", "headline", 400),
    ("summary", "experience_summary", 4000),
    ("clearance", "clearance", 300),
)
_MASTER_EXTRA_SKILL_GROUP = "Additional Skills"
_MASTER_MAX_COMPETENCIES = 40
_MASTER_MAX_SKILLS = 120
_MASTER_MAX_BULLETS = 24
_MASTER_ITEM_MAX_CHARS = 200
_MASTER_BULLET_MAX_CHARS = 1000


def base_document() -> dict:
    """Assemble the Resume Studio editor document for the MASTER resume.

    A straight, model-free mapping of career_db.json into the exact shape the studio's
    renderPreview/applyAction already speak: headline (the profile's optional headline field,
    falling back to the credential line — both candidate-stated, nothing invented), summary,
    competencies, experience (roles with verbatim deliverable bullets), a flattened skills
    list in group order, and clearance. Returns {resume, cover, meta{master: True}}."""
    d = load()
    p = d.get("profile", {}) or {}
    experience = [{
        "title": r.get("title", ""),
        "org": r.get("org", ""),
        "span": f"{r.get('start', '')}–{r.get('end') or 'present'}",
        "bullets": [str(b) for b in (r.get("deliverables") or [])],
    } for r in d.get("roles", []) or []]
    flat_skills = []
    for group in (d.get("skills", {}) or {}).values():
        flat_skills.extend(str(item) for item in (group or {}).get("items", []) or [])
    return {
        "resume": {
            "headline": p.get("headline") or p.get("credential") or "",
            "summary": p.get("experience_summary", ""),
            "competencies": [str(c) for c in d.get("competencies") or []],
            "experience": experience,
            "skills": flat_skills,
            "clearance": p.get("clearance", ""),
        },
        "cover": {},
        "meta": {"master": True, "name": p.get("name", "")},
    }


def _bounded_master_str(value, limit: int, label: str) -> str:
    """Validate one editor text field, refusing (never silently truncating) oversize content."""
    if not isinstance(value, str):
        raise ValueError(f"{label} must be a string")
    text = value.strip()
    if len(text) > limit:
        raise ValueError(f"{label} exceeds {limit} characters")
    return text


def _bounded_master_list(items, max_items: int, max_chars: int, label: str) -> list:
    """Validate one editor string list under explicit count and per-item character bounds."""
    if not isinstance(items, list):
        raise ValueError(f"{label} must be a list")
    if len(items) > max_items:
        raise ValueError(f"{label} exceeds {max_items} items")
    bounded = []
    for item in items:
        text = _bounded_master_str(item, max_chars, f"{label} item")
        if text:
            bounded.append(text)
    return bounded


def _apply_master_texts(raw: dict, resume: dict, changelog: list) -> bool:
    """Apply only the whitelisted text fields the payload actually carries."""
    profile_block = raw.setdefault("profile", {})
    changed = False
    for key, target, limit in _MASTER_TEXT_FIELDS:
        if key not in resume:
            continue
        value = _bounded_master_str(resume.get(key), limit, key)
        if str(profile_block.get(target, "") or "") != value:
            profile_block[target] = value
            changelog.append(f"{key} updated")
            changed = True
    return changed


def _regroup_master_skills(raw: dict, items: list) -> bool:
    """Apply the editor's flat skills list while preserving the profile's group structure.

    Items kept in the list stay in their existing groups, removed items leave their groups,
    and brand-new items land in one dedicated additions group so durable grouping survives."""
    wanted = _low(items)
    groups = raw.setdefault("skills", {})
    changed = False
    covered = set()
    for name in list(groups):
        entry = groups.get(name) or {}
        existing = entry.get("items", []) or []
        kept = [item for item in existing if str(item).strip().lower() in wanted]
        covered.update(str(item).strip().lower() for item in kept)
        if kept != existing:
            entry["items"] = kept
            groups[name] = entry
            changed = True
    additions = [item for item in items if item.lower() not in covered]
    if additions:
        grp = groups.setdefault(_MASTER_EXTRA_SKILL_GROUP, {"items": []})
        grp.setdefault("items", [])
        changed = _append_unique(grp["items"], additions) or changed
    return changed


def _apply_master_lists(raw: dict, resume: dict, changelog: list) -> bool:
    """Apply the whitelisted competencies and skills lists the payload actually carries."""
    changed = False
    if "competencies" in resume:
        items = _bounded_master_list(
            resume.get("competencies"), _MASTER_MAX_COMPETENCIES, _MASTER_ITEM_MAX_CHARS,
            "competencies")
        if items != [str(c) for c in raw.get("competencies") or []]:
            raw["competencies"] = items
            changelog.append(f"competencies updated ({len(items)})")
            changed = True
    if "skills" in resume:
        items = _bounded_master_list(
            resume.get("skills"), _MASTER_MAX_SKILLS, _MASTER_ITEM_MAX_CHARS, "skills")
        if _regroup_master_skills(raw, items):
            changelog.append("skills updated")
            changed = True
    return changed


def _apply_master_bullets(raw: dict, resume: dict, changelog: list) -> bool:
    """Apply per-role bullets, matching each entry by index AND title; any mismatch refuses."""
    if "experience" not in resume:
        return False
    entries = resume.get("experience")
    if not isinstance(entries, list):
        raise ValueError("experience must be a list")
    roles = raw.get("roles", []) or []
    if len(entries) > len(roles):
        raise ValueError("experience has more entries than the profile has roles")
    changed = False
    for index, entry in enumerate(entries):
        if not isinstance(entry, dict):
            raise ValueError(f"experience entry {index + 1} must be an object")
        title = str(entry.get("title", "")).strip()
        role_title = str(roles[index].get("title", "")).strip()
        if title != role_title:
            raise ValueError(
                f"experience entry {index + 1} title does not match the profile role "
                f"(got {title!r}, profile has {role_title!r}) — "
                "master save matches roles by index and title")
        if "bullets" not in entry:
            continue
        bullets = _bounded_master_list(
            entry.get("bullets"), _MASTER_MAX_BULLETS, _MASTER_BULLET_MAX_CHARS,
            f"role {index + 1} bullets")
        if bullets != [str(b) for b in roles[index].get("deliverables") or []]:
            roles[index]["deliverables"] = bullets
            changelog.append(f"role {index + 1} ({title}) bullets updated ({len(bullets)})")
            changed = True
    return changed


def replace_resume_fields(doc: dict) -> dict:
    """Whitelist write-back of the studio's MASTER document into career_db.json.

    Applies ONLY headline, summary, clearance, competencies, skills, and per-role bullets
    (matched by index AND title — any mismatch refuses the whole save with the file
    untouched). Everything else in the payload is ignored. Persistence reuses augment()'s
    rails: pre-image backup with rotation, bounded audit record, atomic replace, and cache
    invalidation. Returns {ok, changed, changelog} or {ok: False, error}."""
    if not isinstance(doc, dict) or not isinstance(doc.get("resume"), dict):
        return {"ok": False, "error": "invalid master resume document"}
    if not config.CAREER_DB.exists():
        return {"ok": False, "error": "no career profile yet — upload a resume first"}
    text = _read_profile_text()
    try:
        raw = json.loads(text)
    except ValueError:
        # A corrupt profile file must answer the refusal shape, not a traceback: the route
        # maps {ok: False} to a readable 422 while a raised exception becomes a generic 500.
        return {"ok": False, "error": "career profile file is unreadable — restore from backup"}
    resume = doc["resume"]
    changelog: list = []
    try:
        changed = _apply_master_bullets(raw, resume, changelog)
        changed = _apply_master_texts(raw, resume, changelog) or changed
        changed = _apply_master_lists(raw, resume, changelog) or changed
        if not changed:
            return _augmentation_result(False, [])
        _persist_augmentation(text, raw, "resume-studio master edit", changelog)
    except ValueError as exc:
        return {"ok": False, "error": str(exc)}
    return _augmentation_result(True, changelog)
